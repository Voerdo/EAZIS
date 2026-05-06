import os
import tempfile
import spacy

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

try:
    import docx2txt
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False

try:
    from docx2python import docx2python
    HAS_DOCX2PYTHON = True
except ImportError:
    HAS_DOCX2PYTHON = False

try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

try:
    import pypandoc
    HAS_PYPANDOC = True
except ImportError:
    HAS_PYPANDOC = False

try:
    import subprocess
    HAS_SUBPROCESS = True
except ImportError:
    HAS_SUBPROCESS = False

DEP_TRANSLATIONS = {
    "ROOT": "Корень", "nsubj": "Подлежащее", "nsubjpass": "Подлежащее (пассив)",
    "dobj": "Прямое дополнение", "obj": "Дополнение", "iobj": "Косвенное дополнение",
    "attr": "Атрибут", "amod": "Определение (прилагательное)", "nummod": "Числовое определение",
    "advmod": "Обстоятельство (наречие)", "advcl": "Придаточное обстоятельство",
    "det": "Артикль/Определитель", "poss": "Притяжательное",
    "xcomp": "Дополнение к сказуемому", "acomp": "Прилагательное сказуемое",
    "npadvmod": "Именное обстоятельство", "pobj": "Объект предлога", "prep": "Предлог",
    "pcomp": "Предложное дополнение", "relcl": "Относительное придаточное",
    "conj": "Союзная связь", "cc": "Сочинительный союз", "mark": "Маркер придаточного",
    "aux": "Вспомогательный глагол", "auxpass": "Вспомогательный (пассив)",
    "neg": "Отрицание", "compound": "Составное слово", "intj": "Междометие",
    "parataxis": "Паратаксис", "appos": "Приложение", "case": "Падеж",
    "ccomp": "Придаточное дополнительное", "acl": "Определительное придаточное",
    "clf": "Классификатор", "fixed": "Фиксированное выражение",
    "flat": "Плоская конструкция", "list": "Список",
    "dislocated": "Вынесенный элемент", "discourse": "Дискурсивный элемент",
    "expl": "Эксплетив", "csubj": "Подлежащее (придаточное)",
    "csubjpass": "Подлежащее придаточное (пассив)", "agent": "Агент",
    "dep": "Неопределённая зависимость", "obl": "Косвенное обстоятельство",
    "nmod": "Именной модификатор", "cop": "Связка", "prt": "Частица фразового глагола",
    "quantmod": "Модификатор количества", "tmod": "Временной модификатор",
    "npmod": "Именной модификатор", "oprd": "Объект предиката",
    "dative": "Дательный падеж", "predet": "Предопределитель",
    "preconj": "Предсоюз", "mwe": "Многословное выражение",
    "goeswith": "Связанное слово", "reparandum": "Исправление",
    "root": "Корень", "punct": "Пунктуация"
}

POS_TRANSLATIONS = {
    "NOUN": "существительное", "VERB": "глагол", "ADJ": "прилагательное",
    "ADV": "наречие", "PRON": "местоимение", "ADP": "предлог",
    "CCONJ": "сочинительный союз", "SCONJ": "подчинительный союз",
    "INTJ": "междометие", "DET": "определитель/артикль", "NUM": "числительное",
    "PROPN": "имя собственное", "AUX": "вспомогательный глагол",
    "PART": "частица", "SYM": "символ", "X": "прочее",
    "PUNCT": "пунктуация", "SPACE": "пробел"
}

try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Модель en_core_web_sm не найдена. Установите: python -m spacy download en_core_web_sm")
    nlp = None


def extract_text_python_docx(filepath):
    if not HAS_PYTHON_DOCX:
        return None
    try:
        doc = Document(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        return "\\n".join(paragraphs)
    except Exception as e:
        print(f"python-docx ошибка: {e}")
        return None

def extract_text_docx2txt(filepath):
    if not HAS_DOCX2TXT:
        return None
    try:
        text = docx2txt.process(filepath)
        return text.strip()
    except Exception as e:
        print(f"docx2txt ошибка: {e}")
        return None

def extract_text_docx2python(filepath):
    if not HAS_DOCX2PYTHON:
        return None
    try:
        with docx2python(filepath) as docx_content:
            text = docx_content.text
        return text.strip()
    except Exception as e:
        print(f"docx2python ошибка: {e}")
        return None

def extract_text_win32com(filepath):
    if not HAS_WIN32COM:
        return None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(os.path.abspath(filepath))
        text = doc.Content.Text
        doc.Close(SaveChanges=False)
        word.Quit()
        return text.strip()
    except Exception as e:
        print(f"win32com ошибка: {e}")
        return None

def extract_text_pypandoc(filepath):
    if not HAS_PYPANDOC:
        return None
    try:
        text = pypandoc.convert_file(filepath, "plain", format="docx")
        return text.strip()
    except Exception as e:
        print(f"pypandoc ошибка: {e}")
        return None

def extract_text_subprocess_catdoc(filepath):
    if not HAS_SUBPROCESS:
        return None
    try:
        result = subprocess.run(
            ["catdoc", filepath],
            capture_output=True, text=True,
            encoding="utf-8", errors="ignore"
        )
        if result.returncode == 0:
            return result.stdout.strip()
        return None
    except Exception as e:
        print(f"catdoc ошибка: {e}")
        return None

def extract_text_subprocess_textutil(filepath):
    if not HAS_SUBPROCESS:
        return None
    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", filepath],
            capture_output=True, text=True,
            encoding="utf-8", errors="ignore"
        )
        if result.returncode == 0:
            return result.stdout.strip()
        return None
    except Exception as e:
        print(f"textutil ошибка: {e}")
        return None

def extract_text_subprocess_libreoffice(filepath):
    if not HAS_SUBPROCESS:
        return None
    try:
        soffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
        if not os.path.exists(soffice_path):
            # Можно попробовать поискать альтернативные пути
            return None
        
        with tempfile.TemporaryDirectory() as tmpdir:
            result = subprocess.run(
                [soffice_path, "--headless", "--convert-to", "txt:Text",
                 "--outdir", tmpdir, filepath],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0:
                txt_file = os.path.join(tmpdir, os.path.splitext(os.path.basename(filepath))[0] + ".txt")
                if os.path.exists(txt_file):
                    with open(txt_file, "r", encoding="utf-8", errors="ignore") as f:
                        return f.read().strip()
        return None
    except Exception as e:
        print(f"libreoffice ошибка: {e}")
        return None

def extract_text_from_doc(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    methods = []
    errors = []

    if ext == ".docx":
        methods = [
            ("python-docx", extract_text_python_docx),
            ("docx2txt", extract_text_docx2txt),
            ("docx2python", extract_text_docx2python),
            ("pypandoc", extract_text_pypandoc),
            ("win32com", extract_text_win32com),
            ("libreoffice", extract_text_subprocess_libreoffice),
        ]
    else:
        methods = [
            ("win32com", extract_text_win32com),
            ("catdoc", extract_text_subprocess_catdoc),
            ("textutil", extract_text_subprocess_textutil),
            ("libreoffice", extract_text_subprocess_libreoffice),
            ("python-docx", extract_text_python_docx),
        ]

    for name, method in methods:
        try:
            text = method(filepath)
            if text and text.strip():
                print(f"Текст успешно извлечен методом: {name}")
                return text, name
        except Exception as e:
            errors.append(f"{name}: {str(e)}")

    error_msg = "; ".join(errors) if errors else "Неизвестная ошибка"
    return None, error_msg


def analyze_sentence(sentence_text):
    if nlp is None:
        return None
    doc = nlp(sentence_text)
    tokens = []
    for i, token in enumerate(doc):
        tokens.append({
            "id": i,
            "text": token.text,
            "lemma": token.lemma_,
            "pos": token.pos_,
            "pos_ru": POS_TRANSLATIONS.get(token.pos_, token.pos_),
            "tag": token.tag_,
            "dep": token.dep_,
            "dep_ru": DEP_TRANSLATIONS.get(token.dep_, token.dep_),
            "head_id": token.head.i if token.head != token else None,
            "head_text": token.head.text if token.head != token else None,
            "is_punct": token.is_punct,
            "is_space": token.is_space
        })

    arcs = []
    for token in doc:
        if token.dep_ != "ROOT" and not token.is_space:
            arcs.append({
                "start": token.head.i,
                "end": token.i,
                "label": token.dep_,
                "label_ru": DEP_TRANSLATIONS.get(token.dep_, token.dep_)
            })

    constituency_tree = None
    try:
        constituency_tree = build_constituency_tree_spacy(doc)
    except Exception as e:
        print(f"Ошибка constituency parsing: {e}")

    return {
        "tokens": tokens,
        "arcs": arcs,
        "constituency_tree": constituency_tree,
        "text": sentence_text
    }

def build_constituency_tree_spacy(doc):
    def build_subtree(token, visited=None):
        if visited is None:
            visited = set()
        if token.i in visited:
            return None
        visited.add(token.i)
        children_trees = []
        for child in token.children:
            if child.i not in visited:
                subtree = build_subtree(child, visited.copy())
                if subtree:
                    children_trees.append(subtree)
        pos_label = token.pos_
        if token.dep_ == "ROOT":
            pos_label = "S"
        elif token.pos_ in ["NOUN", "PROPN", "PRON"]:
            pos_label = "NP"
        elif token.pos_ == "VERB":
            pos_label = "VP"
        elif token.pos_ == "ADP":
            pos_label = "PP"
        elif token.pos_ == "ADJ":
            pos_label = "ADJP"
        elif token.pos_ == "ADV":
            pos_label = "ADVP"
        if children_trees:
            return f"({pos_label} {' '.join(children_trees)} {token.text})"
        return f"({pos_label} {token.text})"

    roots = [token for token in doc if token.dep_ == "ROOT"]
    if roots:
        return build_subtree(roots[0])
    return None

def build_simple_constituency_tree(sentence_text):
    if nlp is None:
        return None
    doc = nlp(sentence_text)
    tokens = list(doc)

    def get_phrase_type(token):
        if token.pos_ in ["NOUN", "PROPN", "PRON", "NUM"]:
            return "NP"
        elif token.pos_ == "VERB":
            return "VP"
        elif token.pos_ == "ADP":
            return "PP"
        elif token.pos_ == "ADJ":
            return "ADJP"
        elif token.pos_ == "ADV":
            return "ADVP"
        elif token.pos_ == "DET":
            return "DT"
        return token.pos_

    def build_tree(token, visited=None):
        if visited is None:
            visited = set()
        if token.i in visited:
            return ""
        visited.add(token.i)
        phrase = get_phrase_type(token)
        children = [c for c in token.children if c.i not in visited]
        if children:
            child_trees = []
            for child in children:
                ct = build_tree(child, visited.copy())
                if ct:
                    child_trees.append(ct)
            if child_trees:
                return f"({phrase} {' '.join(child_trees)} {token.text})"
        return f"({phrase} {token.text})"

    roots = [t for t in tokens if t.dep_ == "ROOT"]
    if roots:
        return build_tree(roots[0])
    return f"(S {' '.join([f'({get_phrase_type(t)} {t.text})' for t in tokens])})"

def get_available_methods():
    return {
        "python-docx": HAS_PYTHON_DOCX,
        "docx2txt": HAS_DOCX2TXT,
        "docx2python": HAS_DOCX2PYTHON,
        "win32com": HAS_WIN32COM,
        "pypandoc": HAS_PYPANDOC,
        "subprocess (catdoc/textutil/libreoffice)": HAS_SUBPROCESS,
    }